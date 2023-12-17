import streamlit as st
from vertexai.language_models import TextGenerationModel
from pydub import AudioSegment
from google.cloud import speech_v1p1beta1 as speech
from google.cloud import aiplatform
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io





def initialize_once():
    if "initialized" not in st.session_state:
        aiplatform.init()
        st.session_state.client = speech.SpeechClient()
        st.session_state.palm_model = TextGenerationModel.from_pretrained("text-bison@001")
        st.session_state.palm_parameters = {
            "temperature": 1,
            "max_output_tokens": 256,
            "top_k": 3,
            "top_p": 0.5
        }        
        st.session_state.initialized = True


def main():
    initialize_once()
    st.title("Meeting Minutes")
    # Widget to upload audio file
    uploaded_file = st.file_uploader("Upload Audio File (Supported Formats: WAV, FLAC, OGG, MP3, M4A)", type=["wav", "flac", "ogg", "mp3", "m4a"])


    # Check if the user uploaded an audio file or recorded audio
    if uploaded_file is not None:
        st.success("Audio file successfully uploaded or recorded!")

        # Additional processing or analysis can be added here

        # Widget to submit the processed audio
        if st.button("Submit"):
            transcript = recognize_audio(uploaded_file)
            st.success("Transcription Generated Successfully")
            summary = meeting_minutes(transcript)
            st.success("Processing complete!")
            doc = write_dict_to_docx(summary)
            # Save the document to a BytesIO object
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Display the download button
            st.download_button(
                label="Download Summary",
                data=doc_io,
                file_name="meeting_minutes.docx",
                key="download_button"
            )
            

def write_dict_to_docx(dictionary):
    document = Document()

    for key, value in dictionary.items():
        # Add key in bold
        key_paragraph = document.add_paragraph()
        key_run = key_paragraph.add_run(key)
        key_run.bold = True
        key_run.font.size = Pt(12)
        key_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add value under the key
        value_paragraph = document.add_paragraph(str(value))
        value_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add a line break between key-value pairs
        document.add_paragraph()

    return document

def generate_text(prompt):
    completion = st.session_state.palm_model.predict(prompt, **st.session_state.palm_parameters)
    return completion.text

def meeting_minutes(transcript):
    abstract_summary = abstract_summary_extraction(transcript)
    key_points = key_points_extraction(transcript)
    action_items = action_item_extraction(transcript)
    keywords = keywords_extraction(transcript)
    sentiment = sentiment_analysis(transcript)
    return {
        'Abstract Summary': abstract_summary,
        'Key Points': key_points,
        'Keywords': keywords,
        'Action Items': action_items,
        'Sentiment': sentiment
    }

def abstract_summary_extraction(transcript):
    instructions = "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
    return generate_text(f"{instructions}\n{transcript}")





def key_points_extraction(transcript):
    instructions = "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
    return generate_text(f"{instructions}\n{transcript}")






def keywords_extraction(transcript):
    instructions = "You are an AI expert in analyzing conversations and extracting keywords. You will be provided with a block of text, and your task is to extract a list of most important keywords from it. Please list the top 10 keywords and use a comma to separate the keywords in the output. "
    return generate_text(f"{instructions}\n{transcript}")



def action_item_extraction(transcript):
    instructions = "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
    return generate_text(f"{instructions}\n{transcript}")



def sentiment_analysis(transcript):
    instructions = "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, neutral, or negative, and provide brief explanations for your analysis where possible."
    return generate_text(f"{instructions}\n{transcript}")


def recognize_audio(input_path, clip_duration=50000):
    audio = AudioSegment.from_file(input_path)
    audio_duration = len(audio) / 1000  # Convert milliseconds to seconds

    num_clips = int((audio_duration // 50) + 1)
    transcript = ""

    # Split the audio into clips
    for i in range(num_clips):
        start_time = i * clip_duration
        end_time = min((i + 1) * clip_duration, audio_duration * 1000)  # Convert seconds to milliseconds
        clip = audio[start_time:end_time]
        audio_bytes = clip.raw_data
        audio_data = speech.RecognitionAudio(content=audio_bytes)
        config = speech.RecognitionConfig(
            encoding=speech.RecognitionConfig.AudioEncoding.LINEAR16,
            sample_rate_hertz=audio.frame_rate,
            language_code="en-US",  # Set the language code accordingly
        )
        response = st.session_state.client.recognize(config=config, audio=audio_data)
        # Print the transcriptions
        for result in response.results:
            transcript += result.alternatives[0].transcript
    return transcript


if __name__ == "__main__":
    main()
